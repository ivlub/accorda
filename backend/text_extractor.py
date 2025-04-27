# This is a pretty simple utility file that extracts text from PDF and DOCX files.
# So it just returns plaintext from the input file.
import os
import logging
from typing import Optional

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
    logging.warning("pypdf library not found. PDF extraction will not work. Run: pip install pypdf")

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx library not found. DOCX extraction will not work. Run: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def _extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extracts text content from a PDF file."""
    if not PdfReader:
        logger.error("pypdf library is required for PDF extraction but not installed.")
        return None
    
    logger.info(f"Attempting to extract text from PDF: {file_path}")
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ""
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n" # Add newline between pages
                except Exception as page_e:
                    logger.warning(f"Could not extract text from page {i} in {file_path}: {page_e}")
            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return text.strip() if text else None
    except FileNotFoundError:
        logger.error(f"PDF file not found: {file_path}")
        return None
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}", exc_info=True)
        return None

def _extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extracts text content from a DOCX file."""
    if not docx:
        logger.error("python-docx library is required for DOCX extraction but not installed.")
        return None
        
    logger.info(f"Attempting to extract text from DOCX: {file_path}")
    try:
        document = docx.Document(file_path)
        text = "\n".join([para.text for para in document.paragraphs if para.text])
        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return text.strip() if text else None
    except FileNotFoundError:
        logger.error(f"DOCX file not found: {file_path}")
        return None
    except Exception as e:
        # python-docx can raise various errors for corrupted files
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}", exc_info=True)
        return None

def extract_text(file_path: str) -> Optional[str]:
    """Extracts text from PDF or DOCX files based on extension."""
    if not os.path.exists(file_path):
        logger.error(f"File does not exist: {file_path}")
        return None
        
    _, extension = os.path.splitext(file_path.lower())
    
    if extension == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif extension == '.docx':
        return _extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type: {extension}. Only .pdf and .docx are supported.")
        return None

# Example Usage (for testing outside the main app)
if __name__ == '__main__':
    # Create dummy files for testing if they don't exist
    # In a real scenario, provide paths to actual files
    TEST_DIR = "test_files"
    os.makedirs(TEST_DIR, exist_ok=True)
    DUMMY_DOCX = os.path.join(TEST_DIR, "dummy.docx")
    DUMMY_PDF = os.path.join(TEST_DIR, "dummy.pdf") # Requires a PDF library to create
    DUMMY_TXT = os.path.join(TEST_DIR, "dummy.txt")

    # Create a dummy docx if python-docx is available and file doesn't exist
    if docx and not os.path.exists(DUMMY_DOCX):
        try:
            doc = docx.Document()
            doc.add_paragraph("This is the first paragraph.")
            doc.add_paragraph("This is the second paragraph with some dummy text.")
            doc.save(DUMMY_DOCX)
            print(f"Created dummy DOCX: {DUMMY_DOCX}")
        except Exception as e:
            print(f"Could not create dummy DOCX: {e}")
    
    # Cannot easily create a valid dummy PDF without more complex dependencies
    # You should manually place a test PDF here if needed
    print(f"Please place a test PDF file at: {DUMMY_PDF} for PDF testing.")

    if not os.path.exists(DUMMY_TXT):
         with open(DUMMY_TXT, "w") as f:
             f.write("This is a text file.")
             print(f"Created dummy TXT: {DUMMY_TXT}")

    print("\n--- Testing Text Extraction ---")
    
    # Test DOCX
    if os.path.exists(DUMMY_DOCX):
        print(f"\nExtracting from: {DUMMY_DOCX}")
        docx_text = extract_text(DUMMY_DOCX)
        if docx_text:
            print("--- DOCX Text ---")
            print(docx_text)
            print("-----------------")
        else:
            print("Failed to extract text from DOCX.")
    else:
        print(f"\nSkipping DOCX test: {DUMMY_DOCX} not found.")

    # Test PDF
    if os.path.exists(DUMMY_PDF):
        print(f"\nExtracting from: {DUMMY_PDF}")
        pdf_text = extract_text(DUMMY_PDF)
        if pdf_text:
            print("--- PDF Text ---")
            print(pdf_text)
            print("--------------")
        else:
            print("Failed to extract text from PDF.")
    else:
        print(f"\nSkipping PDF test: {DUMMY_PDF} not found.")
        
    # Test unsupported file
    print(f"\nExtracting from: {DUMMY_TXT}")
    txt_output = extract_text(DUMMY_TXT)
    if txt_output is None:
        print("Correctly handled unsupported file type (.txt).")
    else:
        print("Error: Did not correctly handle unsupported file type.")
        
    # Test non-existent file
    print(f"\nExtracting from: non_existent_file.pdf")
    non_existent_output = extract_text("non_existent_file.pdf")
    if non_existent_output is None:
        print("Correctly handled non-existent file.")
    else:
        print("Error: Did not correctly handle non-existent file.") 